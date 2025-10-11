"""
Document Processor - Extract and analyze text from documents
"""
import PyPDF2
from docx import Document
import os
from pathlib import Path
from typing import Dict, List

class DocumentProcessor:
    def __init__(self):
        """Initialize document processor"""
        self.supported_formats = ['.pdf', '.docx', '.doc', '.txt']
    
    def extract_text_from_pdf(self, pdf_path: str) -> str:
        """
        Extract text from PDF file
        Args:
            pdf_path: Path to PDF file
        Returns:
            Extracted text as string
        """
        try:
            print(f"📄 Extracting text from PDF: {Path(pdf_path).name}")
            
            text = ""
            with open(pdf_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                num_pages = len(pdf_reader.pages)
                
                print(f"   Total pages: {num_pages}")
                
                for page_num in range(num_pages):
                    page = pdf_reader.pages[page_num]
                    text += page.extract_text() + "\n"
            
            word_count = len(text.split())
            print(f"✅ Extracted {word_count} words from PDF")
            
            return text.strip()
        
        except Exception as e:
            print(f"❌ Error extracting PDF text: {e}")
            raise Exception(f"Failed to extract text from PDF: {e}")
    
    def extract_text_from_docx(self, docx_path: str) -> str:
        """
        Extract text from DOCX file
        Args:
            docx_path: Path to DOCX file
        Returns:
            Extracted text as string
        """
        try:
            print(f"📄 Extracting text from DOCX: {Path(docx_path).name}")
            
            doc = Document(docx_path)
            text = ""
            
            # Extract from paragraphs
            for para in doc.paragraphs:
                text += para.text + "\n"
            
            # Extract from tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        text += cell.text + " "
                text += "\n"
            
            word_count = len(text.split())
            print(f"✅ Extracted {word_count} words from DOCX")
            
            return text.strip()
        
        except Exception as e:
            print(f"❌ Error extracting DOCX text: {e}")
            raise Exception(f"Failed to extract text from DOCX: {e}")
    
    def extract_text_from_txt(self, txt_path: str) -> str:
        """
        Extract text from TXT file
        Args:
            txt_path: Path to TXT file
        Returns:
            Extracted text as string
        """
        try:
            print(f"📄 Reading text file: {Path(txt_path).name}")
            
            with open(txt_path, 'r', encoding='utf-8') as file:
                text = file.read()
            
            word_count = len(text.split())
            print(f"✅ Read {word_count} words from TXT file")
            
            return text.strip()
        
        except Exception as e:
            print(f"❌ Error reading TXT file: {e}")
            raise Exception(f"Failed to read text file: {e}")
    
    def extract_text(self, file_path: str) -> Dict:
        """
        Extract text from document (auto-detect format)
        Args:
            file_path: Path to document
        Returns:
            Dictionary with extracted text and metadata
        """
        if not os.path.exists(file_path):
            raise FileNotFoundError(f"File not found: {file_path}")
        
        file_ext = Path(file_path).suffix.lower()
        
        if file_ext not in self.supported_formats:
            raise ValueError(f"Unsupported file format: {file_ext}")
        
        try:
            # Extract based on file type
            if file_ext == '.pdf':
                text = self.extract_text_from_pdf(file_path)
            elif file_ext in ['.docx', '.doc']:
                text = self.extract_text_from_docx(file_path)
            elif file_ext == '.txt':
                text = self.extract_text_from_txt(file_path)
            else:
                raise ValueError(f"Unsupported format: {file_ext}")
            
            # Calculate metadata
            words = text.split()
            sentences = text.split('.')
            
            return {
                'success': True,
                'text': text,
                'word_count': len(words),
                'character_count': len(text),
                'sentence_count': len([s for s in sentences if s.strip()]),
                'file_type': file_ext,
                'filename': Path(file_path).name
            }
        
        except Exception as e:
            return {
                'success': False,
                'error': str(e),
                'message': f'Failed to extract text from document'
            }
    
    def chunk_text(self, text: str, max_length: int = 512) -> List[str]:
        """
        Split text into chunks for AI analysis
        Args:
            text: Text to split
            max_length: Maximum words per chunk
        Returns:
            List of text chunks
        """
        words = text.split()
        chunks = []
        
        for i in range(0, len(words), max_length):
            chunk = ' '.join(words[i:i + max_length])
            chunks.append(chunk)
        
        print(f"📦 Split text into {len(chunks)} chunks")
        return chunks


# Test function
if __name__ == "__main__":
    processor = DocumentProcessor()
    
    # Test with a file
    test_file = "test.pdf"
    
    if os.path.exists(test_file):
        try:
            result = processor.extract_text(test_file)
            
            if result['success']:
                print(f"\n📊 Document Info:")
                print(f"   Words: {result['word_count']}")
                print(f"   Characters: {result['character_count']}")
                print(f"   Sentences: {result['sentence_count']}")
                print(f"\n📝 Preview: {result['text'][:200]}...")
            else:
                print(f"❌ {result['message']}")
        
        except Exception as e:
            print(f"❌ Test failed: {e}")
    else:
        print(f"⚠️ Test file not found: {test_file}")